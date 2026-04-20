from __future__ import annotations

from pathlib import Path

from src.docx_export import export_aligned_chord_docx


def test_docx_table_layout(sample_aligned: dict, tmp_path: Path):
    from docx import Document

    out = tmp_path / "song.docx"
    export_aligned_chord_docx(
        arrangement=sample_aligned,
        output_path=out,
        title="song",
        transpose_semitones=0,
        capo_fret=0,
    )
    assert out.exists()

    # Reopen with python-docx, assert 2-row tables match the line count and words.
    doc = Document(str(out))
    tables = doc.tables
    expected_lines = len(sample_aligned["lines"])
    assert len(tables) == expected_lines

    for i, t in enumerate(tables):
        assert len(t.rows) == 2
        words = sample_aligned["lines"][i]["words"]
        assert len(t.columns) == len(words)
        for col_idx, w in enumerate(words):
            chord_cell = t.cell(0, col_idx).text
            word_cell = t.cell(1, col_idx).text
            assert word_cell == w["text"]
            if w.get("chord"):
                assert chord_cell == w["chord"]


def test_docx_transpose(sample_aligned: dict, tmp_path: Path):
    from docx import Document

    out = tmp_path / "song_t.docx"
    export_aligned_chord_docx(
        arrangement=sample_aligned,
        output_path=out,
        title="song_t",
        transpose_semitones=2,
        capo_fret=0,
        prefer_flats=True,
    )
    doc = Document(str(out))
    # line 0 had C and G; +2 semitones with flats -> D and A
    first_line_chord_row = doc.tables[0].rows[0].cells
    chords_in_row = [c.text for c in first_line_chord_row if c.text]
    assert "D" in chords_in_row
    assert "A" in chords_in_row


def test_docx_capo_inverts_into_written_chords(sample_aligned: dict, tmp_path: Path):
    from docx import Document

    out = tmp_path / "song_c.docx"
    export_aligned_chord_docx(
        arrangement=sample_aligned,
        output_path=out,
        title="song_c",
        transpose_semitones=0,
        capo_fret=3,  # capo 3: written chords drop 3 semitones
        prefer_flats=True,
    )
    doc = Document(str(out))
    # C - 3 semitones with flats preference -> A
    chord_row = doc.tables[0].rows[0].cells
    chords_in_row = [c.text for c in chord_row if c.text]
    assert "A" in chords_in_row
    # G - 3 -> E
    assert "E" in chords_in_row
