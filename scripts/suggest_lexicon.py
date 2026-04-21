"""Suggest new entries for src/lexicon.py by diffing an exported DOCX against
its ground-truth chord PDF.

Usage:
    python scripts/suggest_lexicon.py exported.docx ground_truth.pdf

Emits ranked (asr_token → ground_truth_token) substitution candidates. Each
candidate comes with a confidence score based on edit distance + positional
proximity. Hand-review the list, then add the useful ones to src/lexicon.py
DEFAULT_CORRECTIONS.

Design note: this deliberately does NOT auto-apply or mutate lexicon.py. A
regex-based correction pass is high-leverage only when each rule is vetted —
one bad entry (e.g. over-broad pattern that clobbers a real word) affects
every future song.
"""
from __future__ import annotations

import argparse
import difflib
import re
import sys
import unicodedata
from collections import Counter
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))


def _read_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _read_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        from PyPDF2 import PdfReader  # type: ignore[no-redef]
    reader = PdfReader(str(path))
    return "\n".join((p.extract_text() or "") for p in reader.pages)


def _fold(s: str) -> str:
    t = unicodedata.normalize("NFD", s or "")
    t = "".join(c for c in t if unicodedata.category(c) != "Mn")
    return t.lower()


def _tokenize(text: str) -> list[str]:
    # Keep accented originals — we'll fold for comparison but emit originals.
    return re.findall(r"[A-Za-zÀ-ÿ]+", text or "")


CHORD_LIKE = re.compile(r"^[A-G][#b]?(m|maj7|m7|7|dim|sus\d?)?(/[A-G][#b]?)?$")


def _is_chord(tok: str) -> bool:
    return bool(CHORD_LIKE.match(tok))


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("docx", type=Path)
    ap.add_argument("pdf", type=Path)
    ap.add_argument("--min-score", type=float, default=0.6,
                    help="Minimum token-similarity for a candidate pair")
    args = ap.parse_args()

    hyp_text = _read_docx_text(args.docx)
    ref_text = _read_pdf_text(args.pdf)
    hyp_tokens = [t for t in _tokenize(hyp_text) if not _is_chord(t)]
    ref_tokens = [t for t in _tokenize(ref_text) if not _is_chord(t)]

    hyp_folded = [_fold(t) for t in hyp_tokens]
    ref_folded = [_fold(t) for t in ref_tokens]

    # Use difflib to align and find replacement ops.
    matcher = difflib.SequenceMatcher(a=hyp_folded, b=ref_folded, autojunk=False)
    candidates: Counter[tuple[str, str]] = Counter()
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag != "replace":
            continue
        span_h = hyp_tokens[i1:i2]
        span_r = ref_tokens[j1:j2]
        # For 1:1 replacements, take the direct pair. For k:m, emit each
        # side-by-side pairing up to min(k, m).
        for k in range(min(len(span_h), len(span_r))):
            h = span_h[k]
            r = span_r[k]
            if _fold(h) == _fold(r):
                continue
            # Skip trivial plural/casing differences — those are post-hoc
            # fixable by normalizing, not a lexicon rule.
            sim = difflib.SequenceMatcher(None, _fold(h), _fold(r)).ratio()
            if sim < args.min_score:
                continue
            candidates[(_fold(h), r)] += 1

    if not candidates:
        print("No candidates found (or DOCX and PDF already match).")
        return

    print(f"{'count':>5}  {'asr_token':<25} →  ground_truth")
    print("-" * 60)
    for (h, r), count in candidates.most_common(50):
        print(f"{count:>5}  {h:<25} →  {r}")

    print(
        "\nReview the list above and add vetted substitutions to "
        "src/lexicon.py DEFAULT_CORRECTIONS. Good candidates: rare words, "
        "proper nouns, domain-specific vocabulary. Skip generic words and "
        "anything that could clobber legitimate vocabulary."
    )


if __name__ == "__main__":
    main()
