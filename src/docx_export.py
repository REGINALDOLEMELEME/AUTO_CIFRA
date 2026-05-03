from __future__ import annotations

import zlib
from pathlib import Path
from typing import Any

from .transpose import effective_semitones, shift_chord


def _set_font(run, name: str, size_pt: int, bold: bool = False) -> None:
    from docx.shared import Pt

    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold


def _set_cell_font(cell, text: str, name: str, size_pt: int, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(text)
    _set_font(run, name=name, size_pt=size_pt, bold=bold)


def _remove_table_borders(table) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl_pr = table._element.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._element.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


def _split_line_on_punctuation(
    words: list[dict[str, Any]], max_words: int = 12
) -> list[list[dict[str, Any]]]:
    if len(words) <= max_words:
        return [words]
    break_chars = (",", ";", ".", ":", "!", "?")
    break_idx = [
        i for i, w in enumerate(words)
        if str(w.get("text", "")).strip().endswith(break_chars)
    ]
    if not break_idx:
        return [words[i:i + max_words] for i in range(0, len(words), max_words)]

    chunks: list[list[dict[str, Any]]] = []
    start = 0
    while start < len(words):
        limit = min(start + max_words, len(words))
        candidates = [i for i in break_idx if start <= i < limit]
        if not candidates and limit == len(words):
            chunks.append(words[start:limit])
            break
        if not candidates:
            chunks.append(words[start:limit])
            start = limit
            continue
        cut = candidates[-1] + 1
        chunks.append(words[start:cut])
        start = cut
    return [c for c in chunks if c]


def _write_chord_word_table(
    document,
    words: list[dict[str, Any]],
    semitones: int,
    prefer_flats: bool,
    body_font: str,
    chord_font: str,
    body_size_pt: int,
    chord_size_pt: int,
) -> None:
    if not words:
        return
    table = document.add_table(rows=2, cols=len(words))
    table.autofit = True
    _remove_table_borders(table)

    for i, w in enumerate(words):
        chord = w.get("chord")
        chord_txt = shift_chord(chord, semitones, prefer_flats=prefer_flats) if chord else ""
        _set_cell_font(
            table.cell(0, i), chord_txt, name=chord_font, size_pt=chord_size_pt, bold=True
        )
        _set_cell_font(
            table.cell(1, i), w.get("text", ""), name=body_font, size_pt=body_size_pt
        )


def export_aligned_chord_docx(
    arrangement: dict[str, Any],
    output_path: Path,
    title: str,
    transpose_semitones: int = 0,
    capo_fret: int = 0,
    prefer_flats: bool = True,
    body_font: str = "Calibri",
    chord_font: str = "Calibri",
    body_size_pt: int = 11,
    chord_size_pt: int = 11,
) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed.") from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    heading = document.add_heading(title, level=1)
    heading.alignment = 0

    semitones = effective_semitones(transpose_semitones, capo_fret)

    meta = document.add_paragraph()
    meta_run = meta.add_run(
        f"Source: {arrangement.get('source_file', '-')}"
        + (f"  |  Transpose: {transpose_semitones:+d}" if transpose_semitones else "")
        + (f"  |  Capo: {capo_fret}" if capo_fret else "")
    )
    _set_font(meta_run, name=body_font, size_pt=10)
    meta.paragraph_format.space_after = Pt(6)

    warnings = arrangement.get("warnings", []) or []
    if warnings:
        warn_p = document.add_paragraph()
        warn_h = warn_p.add_run("Notes:")
        _set_font(warn_h, name=body_font, size_pt=10, bold=True)
        for warning in warnings:
            wp = document.add_paragraph()
            _set_font(wp.add_run(f"- {warning}"), name=body_font, size_pt=10)

    current_section: str | None = None
    for line in arrangement.get("lines", []) or []:
        section = line.get("section")
        if section and section != current_section:
            current_section = section
            sh = document.add_paragraph()
            sh.paragraph_format.space_before = Pt(10)
            sh.paragraph_format.space_after = Pt(2)
            _set_font(sh.add_run(f"[{section}]"), name=body_font, size_pt=12, bold=True)

        words = line.get("words") or []
        if not words:
            text = str(line.get("lyric_line") or "").strip()
            if text:
                p = document.add_paragraph()
                _set_font(p.add_run(text), name=body_font, size_pt=body_size_pt)
            continue

        for chunk in _split_line_on_punctuation(words, max_words=12):
            _write_chord_word_table(
                document=document,
                words=chunk,
                semitones=semitones,
                prefer_flats=prefer_flats,
                body_font=body_font,
                chord_font=chord_font,
                body_size_pt=body_size_pt,
                chord_size_pt=chord_size_pt,
            )
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)

    document.save(str(output_path))
    return output_path


def export_transcription_docx(transcription: dict[str, Any], output_path: Path, title: str) -> Path:
    """Simple lyrics-only export used by the non-arrangement endpoint."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed.") from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(f"Source: {transcription.get('source_file', '-')}")

    for segment in transcription.get("segments", []):
        text = str(segment.get("text", "")).strip()
        if not text:
            continue
        start = float(segment.get("start", 0.0) or 0.0)
        document.add_paragraph(f"[{start:0.2f}] {text}")

    document.save(str(output_path))
    return output_path


def _pdf_escape(text: str) -> str:
    return (
        text.replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
    )


def _pdf_text_width(text: str, font_size: float) -> float:
    return len(text) * font_size * 0.60


def _pdf_token_columns(
    words: list[dict[str, Any]],
    semitones: int,
    prefer_flats: bool,
) -> tuple[str, str]:
    chord_tokens: list[str] = []
    lyric_tokens: list[str] = []
    for w in words:
        lyric = str(w.get("text") or "")
        chord = str(w.get("chord") or "")
        chord_txt = shift_chord(chord, semitones, prefer_flats=prefer_flats) if chord else ""
        width = max(len(lyric), len(chord_txt), 1)
        chord_tokens.append(chord_txt.ljust(width))
        lyric_tokens.append(lyric.ljust(width))
    return " ".join(chord_tokens).rstrip(), " ".join(lyric_tokens).rstrip()


def _pdf_lines_for_arrangement(
    arrangement: dict[str, Any],
    title: str,
    transpose_semitones: int,
    capo_fret: int,
    prefer_flats: bool,
) -> list[tuple[str, str]]:
    semitones = effective_semitones(transpose_semitones, capo_fret)
    rows: list[tuple[str, str]] = [
        ("title", title),
        (
            "meta",
            f"Source: {arrangement.get('source_file', '-')}"
            + (f"  |  Transpose: {transpose_semitones:+d}" if transpose_semitones else "")
            + (f"  |  Capo: {capo_fret}" if capo_fret else ""),
        ),
    ]
    warnings = arrangement.get("warnings", []) or []
    if warnings:
        rows.append(("warn_head", "Notes:"))
        for warning in warnings:
            rows.append(("warn", f"- {warning}"))

    current_section: str | None = None
    for line in arrangement.get("lines", []) or []:
        section = line.get("section")
        if section and section != current_section:
            current_section = section
            rows.append(("gap_sm", ""))
            rows.append(("section", f"[{section}]"))
        words = line.get("words") or []
        if not words:
            text = str(line.get("lyric_line") or "").strip()
            if text:
                rows.append(("body", text))
                rows.append(("gap_sm", ""))
            continue
        for chunk in _split_line_on_punctuation(words, max_words=12):
            chord_line, lyric_line = _pdf_token_columns(chunk, semitones, prefer_flats)
            rows.append(("chord", chord_line))
            rows.append(("body_mono", lyric_line))
            rows.append(("gap_sm", ""))
    return rows


def _pdf_content_streams(rows: list[tuple[str, str]]) -> list[bytes]:
    page_w = 595
    page_h = 842
    margin_x = 48
    margin_y = 52
    title_size = 18
    meta_size = 10
    section_size = 12
    chord_size = 11
    body_size = 11
    leading = {
        "title": 24,
        "meta": 16,
        "warn_head": 14,
        "warn": 13,
        "section": 16,
        "chord": 13,
        "body": 14,
        "body_mono": 14,
        "gap_sm": 8,
    }

    pages: list[list[str]] = []
    current_page: list[str] = []
    y = page_h - margin_y

    def start_page() -> list[str]:
        return ["BT", f"1 0 0 1 {margin_x} {page_h - margin_y} Tm"]

    current_page = start_page()

    for kind, text in rows:
        step = leading[kind]
        if y - step < margin_y:
            current_page.append("ET")
            pages.append(current_page)
            y = page_h - margin_y
            current_page = start_page()
        y -= step
        if kind == "gap_sm":
            current_page.append(f"1 0 0 1 {margin_x} {y:.2f} Tm")
            continue
        if kind == "title":
            font = "F2"
            size = title_size
        elif kind == "meta":
            font = "F1"
            size = meta_size
        elif kind in {"section", "warn_head"}:
            font = "F2"
            size = section_size if kind == "section" else meta_size
        elif kind == "warn":
            font = "F1"
            size = meta_size
        else:
            font = "F3"
            size = chord_size if kind == "chord" else body_size
        current_page.append(f"1 0 0 1 {margin_x} {y:.2f} Tm")
        current_page.append(f"/{font} {size} Tf")
        current_page.append(f"({_pdf_escape(text)}) Tj")

    current_page.append("ET")
    pages.append(current_page)
    return [
        "\n".join(page).encode("latin-1", errors="replace")
        for page in pages
    ]


def _pdf_obj(obj_id: int, payload: bytes) -> bytes:
    return f"{obj_id} 0 obj\n".encode("ascii") + payload + b"\nendobj\n"


def export_aligned_chord_pdf(
    arrangement: dict[str, Any],
    output_path: Path,
    title: str,
    transpose_semitones: int = 0,
    capo_fret: int = 0,
    prefer_flats: bool = True,
    body_font: str = "Calibri",
    chord_font: str = "Calibri",
    body_size_pt: int = 11,
    chord_size_pt: int = 11,
) -> Path:
    del body_font, chord_font, body_size_pt, chord_size_pt
    output_path.parent.mkdir(parents=True, exist_ok=True)

    rows = _pdf_lines_for_arrangement(
        arrangement=arrangement,
        title=title,
        transpose_semitones=transpose_semitones,
        capo_fret=capo_fret,
        prefer_flats=prefer_flats,
    )
    raw_streams = _pdf_content_streams(rows)

    objects: list[bytes] = []
    objects.append(_pdf_obj(1, b"<< /Type /Catalog /Pages 2 0 R >>"))

    page_count = max(1, len(raw_streams))
    font_obj_start = 3 + page_count * 2
    page_obj_ids = [3 + i * 2 for i in range(page_count)]
    stream_obj_ids = [4 + i * 2 for i in range(page_count)]
    kids = " ".join(f"{obj_id} 0 R" for obj_id in page_obj_ids)
    objects.append(_pdf_obj(2, f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>".encode("ascii")))

    for page_obj_id, stream_obj_id, raw_stream in zip(page_obj_ids, stream_obj_ids, raw_streams):
        page_payload = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
            f"/Resources << /Font << /F1 {font_obj_start} 0 R /F2 {font_obj_start + 1} 0 R /F3 {font_obj_start + 2} 0 R >> >> "
            f"/Contents {stream_obj_id} 0 R >>"
        ).encode("ascii")
        objects.append(_pdf_obj(page_obj_id, page_payload))
        compressed = zlib.compress(raw_stream)
        stream_payload = (
            f"<< /Length {len(compressed)} /Filter /FlateDecode >>\nstream\n".encode("ascii")
            + compressed
            + b"\nendstream"
        )
        objects.append(_pdf_obj(stream_obj_id, stream_payload))

    objects.append(_pdf_obj(font_obj_start, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"))
    objects.append(_pdf_obj(font_obj_start + 1, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>"))
    objects.append(_pdf_obj(font_obj_start + 2, b"<< /Type /Font /Subtype /Type1 /BaseFont /Courier >>"))

    pdf = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj)
    xref_start = len(pdf)
    pdf.extend(f"xref\n0 {len(offsets)}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        pdf.extend(f"{off:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        (
            f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\n"
            f"startxref\n{xref_start}\n%%EOF\n"
        ).encode("ascii")
    )
    output_path.write_bytes(bytes(pdf))
    return output_path
